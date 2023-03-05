import docx
from docx import Document

doc = Document("fun.docx")


for i in range(1,len(doc.paragraphs)):
    doc.paragraphs[i].Alignment = 1
    print(doc.paragraphs[i].text)


# Getting the Full Text from a .docx File




# Styling Paragraph And Run Object
print(doc.paragraphs[0].runs[0].text)


# Run Attributes




# WRITING WORD DOCUMENTS
doc.add_paragraph("hello world")
doc.save("fun.docx")




doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)

doc.add_picture("before.jpg")

doc.add_paragraph("""This creates a two-page Word document with This is on the first page! on the first page and This is
on the second page! on the second. Even though there was still plenty of space on the first page after
the text This is on the first page!, we forced the next paragraph to begin on a new page by inserting a
page break after the first run of the first paragraph""")
doc.paragraphs[1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('headings.docx')


